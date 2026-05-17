import os
import sys
import csv
import logging
import datetime
import re
import sqlite3
from pathlib import Path
from html import escape
from flask import Flask, jsonify, render_template, request, send_from_directory
from werkzeug.utils import secure_filename
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
from bs4 import BeautifulSoup

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    stream=sys.stdout
)
logger = logging.getLogger(__name__)

app = Flask(__name__)

# Paths
BASE_DIR = Path(__file__).parent
DATA_DIR = BASE_DIR / "data"
CSV_PATH  = DATA_DIR / "pdf_data.csv"
DOCX_PATH = DATA_DIR / "pdf_report.docx"
DB_PATH   = DATA_DIR / "products.db"


# ============================================================
# Shared helpers
# ============================================================

def ensure_data_directory():
    """Create data directory if it doesn't exist."""
    if not DATA_DIR.exists():
        DATA_DIR.mkdir(parents=True, exist_ok=True)
        logger.info(f"Created data directory at: {DATA_DIR}")
        return True
    return False


def get_db_connection():
    """Return a cached SQLite connection (one per thread via Flask)."""
    ensure_data_directory()
    conn = sqlite3.connect(str(DB_PATH))
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    """Create the products table if it does not yet exist."""
    conn = get_db_connection()
    conn.execute("""
        CREATE TABLE IF NOT EXISTS products (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            title       TEXT,
            price       TEXT,
            rating      TEXT,
            comments_count TEXT,
            image_url   TEXT,
            product_url TEXT,
            created_at  TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    conn.commit()
    conn.close()
    logger.info("Database initialised.")


# --- PDF helpers (kept — do not break anything that still calls them) ---

def extract_metadata_from_text(text):
    """Attempt to extract year and journal from PDF text."""
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    title = year = journal = abstract = ""
    if not lines:
        return title, year, journal, abstract
    title = lines[0]
    if len(lines) > 1 and len(lines[1]) < 100:
        title += " " + lines[1]
    title = title[:100].strip()
    year_pattern = r'\b(19[0-9]{2}|20[0-9]{2})\b'
    import re
    for line in lines[:10]:
        match = re.search(year_pattern, line)
        if match:
            year = match.group(0)
            break
    journal_keywords = ['journal', 'conference', 'proceedings', 'arxiv', 'volume', 'pp.']
    for line in lines[:15]:
        if any(k in line.lower() for k in journal_keywords):
            journal = line[:100].strip()
            break
    stop_keywords = ['keywords', 'introduction', 'references', 'related work', 'methodology']
    abstract_lines = []
    for i, line in enumerate(lines):
        if i == 0:
            continue
        if any(k in line.lower() for k in stop_keywords):
            break
        if len(line) > 20:
            abstract_lines.append(line)
        if len(' '.join(abstract_lines)) > 500:
            break
    abstract = ' '.join(abstract_lines)[:300].strip()
    return title, year, journal, abstract


def scan_pdfs(force_rescan=False):
    """Scan PDF files in data directory and generate CSV index."""
    ensure_data_directory()
    if not force_rescan and CSV_PATH.exists():
        logger.info(f"CSV already exists at {CSV_PATH}. Use force=True to rescan.")
        return True, "CSV already exists", count_pdfs_in_csv()
    pdf_files = sorted(DATA_DIR.glob("*.pdf"))
    if not pdf_files:
        logger.warning("No PDF files found in data/ folder")
        return False, "No PDF files found in data/ folder", 0
    logger.info(f"Scanning {len(pdf_files)} PDF file(s)...")
    data = []
    for pdf_file in pdf_files:
        try:
            with open(pdf_file, "rb") as f:
                reader = PdfReader(f)
                text = reader.pages[0].extract_text() or "" if reader.pages else ""
                title, year, journal, abstract = extract_metadata_from_text(text)
                if not title:
                    title = pdf_file.stem.replace('_', ' ').title()
                if not abstract and text:
                    abstract = text[:300].strip()
                data.append({
                    "filename": pdf_file.name,
                    "title": title,
                    "year": year,
                    "journal": journal,
                    "abstract": abstract
                })
                logger.info(f"  [OK] Processed: {pdf_file.name}")
        except Exception as e:
            logger.error(f"  [ERROR] Error reading {pdf_file.name}: {e}")
            data.append({
                "filename": pdf_file.name,
                "title": f"[ERROR] {pdf_file.name}",
                "year": "", "journal": "", "abstract": f"Error: {e}"
            })
    fieldnames = ["filename", "title", "year", "journal", "abstract"]
    try:
        with open(CSV_PATH, "w", newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(data)
        logger.info(f"[SUCCESS] CSV generated: {CSV_PATH} ({len(data)} entries)")
        try:
            generate_docx_report(data)
        except Exception as e:
            logger.error(f"[WARN] CSV saved but Word report generation failed: {e}")
        return True, f"Successfully indexed {len(data)} PDF(s)", len(data)
    except Exception as e:
        logger.error(f"Failed to write CSV: {e}")
        return False, f"Failed to write CSV: {e}", 0


def count_pdfs_in_csv():
    """Count how many PDFs are in the CSV."""
    if not CSV_PATH.exists():
        return 0
    try:
        with open(CSV_PATH, "r", encoding="utf-8") as f:
            return sum(1 for _ in f) - 1
    except:
        return 0


def generate_docx_report(data):
    """Build a clean Word document from indexed PDF data."""
    ensure_data_directory()
    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Inches(1.0)
    section.top_margin  = section.bottom_margin = Inches(1.0)
    title_para = doc.add_heading("PDF Search Index Report", level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Generated: {now}  |  Total PDFs indexed: {len(data)}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()
    def add_entry(entry, index):
        if index > 0:
            doc.add_page_break()
        filename = entry.get("filename", "Unknown")
        h = doc.add_heading(f"Entry {index + 1}: {filename}", level=2)
        h.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x7A)
        def add_label_value(label, value):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r1 = p.add_run(f"{label}  ")
            r1.bold = True
            r1.font.size = Pt(11)
            r1.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            r2 = p.add_run(value or "N/A")
            r2.font.size = Pt(11)
        add_label_value("Filename:", filename)
        add_label_value("Title:",    entry.get("title", "N/A"))
        add_label_value("Year:",     entry.get("year",  "N/A"))
        add_label_value("Journal:",  entry.get("journal", "N/A"))
        abstract = entry.get("abstract", "N/A") or "N/A"
        abs_para = doc.add_paragraph()
        abs_para.paragraph_format.space_before = Pt(2)
        abs_para.paragraph_format.space_after  = Pt(2)
        abs_label = abs_para.add_run("Abstract:  ")
        abs_label.bold = True
        abs_label.font.size = Pt(11)
        abs_label.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        abs_para.add_run(abstract).font.size = Pt(11)
        sep = doc.add_paragraph("─" * 80)
        sep.runs[0].font.size  = Pt(8)
        sep.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    for idx, entry in enumerate(data):
        add_entry(entry, idx)
    doc.save(str(DOCX_PATH))
    logger.info(f"Word report generated: {DOCX_PATH} ({len(data)} entries)")
    return DOCX_PATH


def get_search_results(query):
    """Search the CSV index for matching PDFs."""
    if not CSV_PATH.exists():
        logger.warning("CSV not found. Cannot search.")
        return []
    results = []
    try:
        with open(CSV_PATH, "r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                searchable_text = (
                    (row.get("title", "") or "") + " " +
                    (row.get("abstract", "") or "") + " " +
                    (row.get("journal", "") or "") + " " +
                    (row.get("year", "") or "")
                ).lower()
                if query.lower() in searchable_text:
                    results.append(row)
    except Exception as e:
        logger.error(f"Error reading CSV: {e}")
        return []
    return results


# ============================================================
# Digikala Crawler
# ============================================================

SESSION = requests.Session()
SESSION.headers.update({
    "User-Agent":
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/120.0.0.0 Safari/537.36",
    "Accept-Language": "en-US,en;q=0.9,fa;q=0.8",
    "Accept":
        "text/html,application/xhtml+xml,application/xml;"
        "q=0.9,*/*;q=0.8",
})


def _text(el):
    """Strip whitespace from a BeautifulSoup element's text or return None."""
    t = el.get_text(" ", strip=True) if el else None
    return t if t else None


def _attr(el, name):
    """Safely read an attribute from a BeautifulSoup element."""
    return el.get(name) if el else None


def crawl_digikala_product(url):
    """
    Scrape a single Digikala product page.

    Returns a dict with keys:
        title, price, rating, comments_count, image_url, product_url

    All individual fields may be None on failure — the calling code
    decides what to do with partial results.
    """
    logger.info(f"Crawling: {url}")

    try:
        response = SESSION.get(url, timeout=15)
        response.raise_for_status()
        logger.info(f"HTTP {response.status_code} — page downloaded ({len(response.content)} bytes)")
    except requests.exceptions.Timeout:
        logger.error(f"Timeout crawling {url}")
        return None
    except requests.exceptions.HTTPError as e:
        logger.error(f"HTTP error {e.response.status_code} for {url}")
        return None
    except requests.exceptions.RequestException as e:
        logger.error(f"Request error for {url}: {e}")
        return None

    soup = BeautifulSoup(response.text, "html.parser")

    # ---------- title ----------
    # Digikala uses <h1 class="text-h3 …"> for the product title
    title = (_text(soup.find("h1"))
             or _text(soup.find("h1", class_=re.compile("text-h3|product-title|pdp-title", re.I)))
             or _text(soup.find("meta", property="og:title"))
             or _text(soup.find("title")))

    # ---------- price ----------
    # Primary: <span class="text-dark-green font-bold text-h3 price-value">
    price_el = (
        soup.find("span", class_=re.compile("price-value|text-h3.*price", re.I))
        or soup.find(attrs={"data-testid": "buy-box-price"})
    )
    price = _text(price_el)

    # ---------- rating ----------
    # e.g. "4.5" from <span class="text-body-3 rating-count">
    rating_el = (
        soup.find("span", class_=re.compile("rating-count|text-body-3", re.I))
    )
    rating = _text(rating_el)

    # ---------- comments count ----------
    comments_el = (
        soup.find("a", href=re.compile("comment", re.I))
        or soup.find(["span", "a"], class_=re.compile("comment|review-count", re.I))
    )
    comments_count = _text(comments_el)

    # ---------- main image ----------
    # og:image meta tag is the most reliable source
    image_el = soup.find("img", class_=re.compile("pdp-image|main-image|product-image", re.I))
    image_url = None
    if image_el:
        for attr in ("data-src", "data-lazy-src", "src"):
            src = image_el.get(attr)
            if src and not src.startswith("data:"):
                image_url = src
                break
    if not image_url:
        og_image = soup.find("meta", property="og:image")
        if og_image:
            image_url = og_image.get("content")

    result = {
        "title":         title,
        "price":         price,
        "rating":        rating,
        "comments_count":comments_count,
        "image_url":     image_url,
        "product_url":   url,
    }

    logger.info(
        f"Crawl result — title={title[:60] if title else 'N/A'} | "
        f"price={price or 'N/A'} | rating={rating or 'N/A'}"
    )
    return result


# ============================================================
# Flask Routes
# ============================================================

@app.route("/")
def index():
    """Serve the main UI."""
    try:
        return render_template("index.html")
    except Exception as e:
        logger.error(f"Template error: {e}")
        return jsonify({"error": "Template not found", "message": str(e)}), 500


@app.route("/health")
def health():
    """Health check endpoint."""
    return jsonify({
        "status": "OK",
        "port": 5001
    })


# --- PDF routes (kept for backward-compatibility; UI no longer calls them) ---

@app.route("/scan", methods=["GET", "POST"])
def scan():
    """Trigger PDF scanning and CSV generation."""
    if request.method == "GET":
        if CSV_PATH.exists():
            count = count_pdfs_in_csv()
            return jsonify({
                "status": "already_scanned",
                "message": f"CSV already exists with {count} entries",
                "count": count
            })
        else:
            success, message, count = scan_pdfs(force_rescan=True)
            return jsonify({
                "status": "success" if success else "no_pdfs",
                "message": message,
                "count": count
            })
    else:
        force = request.args.get("force", "false").lower() == "true"
        success, message, count = scan_pdfs(force_rescan=force)
        return jsonify({
            "status": "success" if success else "error",
            "message": message,
            "count": count
        })


@app.route("/search")
def search():
    """Search the CSV index for matching PDFs."""
    query = request.args.get("q", "").strip()
    if not query:
        return jsonify({"results": [], "message": "Empty query"})
    logger.info(f"Searching for: {query}")
    results = get_search_results(query)
    logger.info(f"Found {len(results)} results")
    return jsonify({"query": query, "count": len(results), "results": results})


@app.route("/upload", methods=["POST"])
def upload():
    """Accept a PDF file upload, save it to DATA_DIR, and re-index."""
    ensure_data_directory()
    if "pdf" not in request.files:
        logger.warning("Upload request received with no file field 'pdf'.")
        return jsonify({"status": "error", "message": "No file field 'pdf' in request."}), 400
    uploaded_file = request.files["pdf"]
    if uploaded_file.filename == "" or uploaded_file.filename is None:
        logger.warning("Upload request received with empty filename.")
        return jsonify({"status": "error", "message": "No file selected."}), 400
    filename = secure_filename(uploaded_file.filename)
    if not filename.lower().endswith(".pdf"):
        logger.warning(f"Upload rejected — not a PDF: {filename}")
        return jsonify({"status": "error", "message": "Only PDF files are allowed."}), 400
    save_path = DATA_DIR / filename
    if save_path.exists():
        base, ext = save_path.stem, save_path.suffix
        counter = 1
        while save_path.exists():
            save_path = DATA_DIR / f"{base}_{counter}{ext}"
            counter += 1
        logger.info(f"Duplicate filename detected. Saving as: {save_path.name}")
    else:
        logger.info(f"Saving upload as: {filename}")
    try:
        uploaded_file.save(str(save_path))
        logger.info(f"Upload saved successfully: {save_path.name}")
    except Exception as e:
        logger.error(f"Failed to save uploaded file: {e}")
        return jsonify({"status": "error", "message": f"Failed to save file: {e}"}), 500
    logger.info("Triggering PDF re-index after upload...")
    success, message, count = scan_pdfs(force_rescan=True)
    if success:
        logger.info(f"Upload + index completed. Indexed {count} PDF(s).")
        return jsonify({"status": "success", "message": f"Uploaded successfully. Indexed {count} PDFs."})
    else:
        logger.error(f"Upload saved but indexing failed: {message}")
        return jsonify({"status": "error", "message": f"File uploaded but indexing failed: {message}"}), 500


@app.route("/debug/files")
def debug_files():
    """Debug endpoint to show filesystem state."""
    import os
    try:
        cwd   = os.getcwd()
        base  = str(BASE_DIR.absolute())
        data  = str(DATA_DIR.absolute())
        exists = DATA_DIR.exists()
        files = os.listdir(DATA_DIR) if exists else []
        pdfs   = [str(f) for f in DATA_DIR.glob("*.pdf")] if exists else []
        detail = []
        if exists:
            for f in os.listdir(DATA_DIR):
                fp = DATA_DIR / f
                try:
                    detail.append({"name": f, "is_file": fp.is_file(),
                                   "size": fp.stat().st_size, "ext": fp.suffix.lower()})
                except Exception as e:
                    detail.append({"name": f, "error": str(e)})
        return jsonify({"cwd": cwd, "base_dir": base, "data_dir": data,
                        "exists": exists, "files": files, "pdfs": pdfs, "detail": detail})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/status")
def api_status():
    """Return app status."""
    conn = get_db_connection()
    product_count = conn.execute("SELECT COUNT(*) FROM products").fetchone()[0]
    conn.close()
    return jsonify({
        "data_directory":       str(DATA_DIR),
        "data_directory_exists": DATA_DIR.exists(),
        "crawler_ready":         True,
        "product_count":         product_count,
        "message": (f"{product_count} product(s) crawled"
                    if product_count else "No products crawled yet. Add a Digikala URL above.")
    })


# ======== NEW: Digikala Crawler route ========

@app.route("/crawl", methods=["POST"])
def crawl():
    """
    Receive a Digikala product URL, scrape it, save to SQLite.
    JSON body: {"url": "https://digikala.com/product/..."}
    """
    body    = request.get_json(silent=True) or {}
    raw_url = (body.get("url") or "").strip()

    # --- basic validation ---
    if "digikala.com" not in raw_url.lower():
        logger.warning(f"Crawl rejected — not a digikala.com URL: {raw_url}")
        return jsonify({"status": "error",
                        "message": "URL must be a valid Digikala product page."}), 400

    # normalize URL
    if not raw_url.startswith("http"):
        raw_url = "https://" + raw_url

    logger.info(f" crawl request received — {raw_url}")

    info = crawl_digikala_product(raw_url)

    if info is None:
        logger.error(f"Crawl failed for {raw_url} — no data returned.")
        return jsonify({"status": "error",
                        "message": "Failed to crawl product. The page may be unavailable or blocked."}), 500

    title  = info.get("title")     or "Unknown"
    price  = info.get("price")     or "—"
    rating = info.get("rating")    or "—"
    comments = info.get("comments_count") or "—"
    img    = info.get("image_url") or ""
    purl   = info.get("product_url") or raw_url

    try:
        conn = get_db_connection()
        conn.execute("""
            INSERT INTO products (title, price, rating, comments_count, image_url, product_url)
            VALUES (?, ?, ?, ?, ?, ?)
        """, (title, price, rating, comments, img, purl))
        conn.commit()
        conn.close()
        logger.info(f"Product saved to DB: {title[:80]}")
    except Exception as e:
        logger.error(f"Database insert failed: {e}")
        return jsonify({"status": "error", "message": f"Saved crawl data but DB write failed: {e}"}), 500

    return jsonify({
        "status": "success",
        "message": f"Product crawled successfully: {title}",
        "product": info
    })


# ======== NEW: Products API ========

@app.route("/products")
def list_products():
    """Return all crawled products from SQLite (newest first)."""
    try:
        conn = get_db_connection()
        rows = conn.execute(
            "SELECT * FROM products ORDER BY id DESC"
        ).fetchall()
        conn.close()
        products = [dict(row) for row in rows]
        return jsonify({"count": len(products), "products": products})
    except Exception as e:
        logger.error(f"Products API error: {e}")
        return jsonify({"error": str(e)}), 500


# ======== PDF Word-report download (kept) ========

@app.route("/download")
def download_docx():
    """Download the generated Word report."""
    if DOCX_PATH.exists():
        return send_from_directory(DATA_DIR, "pdf_report.docx", as_attachment=True)
    else:
        return jsonify({"error": "Word report not found. Run /scan first."}), 404


@app.route("/download-docx")
def download_docx_legacy():
    """Legacy alias for /download."""
    return download_docx()


# ============================================================
# App Initialization
# ============================================================

def initialize_app():
    """Initialize the application on startup (both local and production)."""
    logger.info("=" * 60)
    logger.info("DIGIKALA PRODUCT CRAWLER initializing...")
    logger.info(f"BASE_DIR: {BASE_DIR.absolute()}")
    logger.info(f"DATA_DIR: {DATA_DIR.absolute()}")
    logger.info("=" * 60)

    ensure_data_directory()
    init_db()

    logger.info("=" * 60)
    logger.info("Initialization complete. Ready to serve requests.")

    template_path = BASE_DIR / "templates" / "index.html"
    if not template_path.exists():
        logger.error(f"FATAL: Template missing: {template_path}")
        sys.exit(1)


# ============================================================
# App Entry Point
# ============================================================

if __name__ == "__main__":
    initialize_app()
    port = int(os.environ.get("PORT", 5001))
    logger.info(f"Starting Flask development server on http://0.0.0.0:{port}")
    app.run(host="0.0.0.0", port=port, debug=False, use_reloader=False)
else:
    # Production (gunicorn)
    initialize_app()
